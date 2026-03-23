"""
build_manual_docx.py
--------------------
Reads MANUAL.txt and produces Isotope_Dashboard_Manual.docx using the
Curium house style (Calibri, brand colours, logo header).

Run:  python build_manual_docx.py
"""

import io
import re
import textwrap
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml

# ── Brand colours ────────────────────────────────────────────────────────────
PURPLE     = RGBColor(0x4B, 0x1E, 0x52)   # Curium dark purple
PINK       = RGBColor(0xE5, 0x06, 0x95)   # Curium hot pink
GRAY_TEXT  = RGBColor(0x80, 0x80, 0x80)   # Tagline / footer grey
BLACK      = RGBColor(0x00, 0x00, 0x00)
CODE_BG    = RGBColor(0xF2, 0xF2, 0xF2)

FONT_NAME  = "Calibri"
BODY_SIZE  = Pt(10)
H1_SIZE    = Pt(16)
H2_SIZE    = Pt(13)
H3_SIZE    = Pt(11)
CODE_SIZE  = Pt(9)


# ── Logo loader ───────────────────────────────────────────────────────────────
def _build_logo_png() -> bytes:
    """Return the official Curium logo PNG bytes from logo.png in the repo root."""
    logo_path = Path(__file__).parent / "logo.png"
    return logo_path.read_bytes()


# ── docx helpers ──────────────────────────────────────────────────────────────
def _set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False,
              italic=False, color=BLACK, underline=False):
    run.font.name      = name
    run.font.size      = size
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color
    run.font.underline = underline


def _para_spacing(para, before=Pt(0), after=Pt(4), line=None):
    pPr = para._p.get_or_add_pPr()
    spg = OxmlElement("w:spacing")
    spg.set(qn("w:before"), str(int(before.pt * 20)))
    spg.set(qn("w:after"),  str(int(after.pt  * 20)))
    if line is not None:
        spg.set(qn("w:line"),     str(int(line.pt * 20)))
        spg.set(qn("w:lineRule"), "exact")
    pPr.append(spg)


def _shade_cell(cell, color_hex: str):
    """Apply a solid background shade to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)


def _add_page_border(doc):
    """Add a thin purple top border on every page via sectPr."""
    # (skipped for simplicity — Word page borders are complex XML)
    pass


def _set_page_margins(doc):
    for section in doc.sections:
        section.page_width   = Cm(21)
        section.page_height  = Cm(29.7)
        section.left_margin  = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.top_margin   = Cm(2.5)
        section.bottom_margin = Cm(2.0)


def _add_header(doc, logo_bytes: bytes):
    """Add Curium logo + document title to the first-page header."""
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header  = section.first_page_header

    # Clear any default paragraph
    for p in header.paragraphs:
        p.clear()

    # Logo in a 1×2 table so text can sit alongside it
    tbl = header.add_table(1, 2, width=Cm(16))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"
    # Remove table borders
    tbl_elem = tbl._tbl
    tblPr    = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    # Left cell: logo
    cell_logo = tbl.cell(0, 0)
    cell_logo.width = Cm(6)
    p_logo = cell_logo.paragraphs[0]
    run_logo = p_logo.add_run()
    run_logo.add_picture(io.BytesIO(logo_bytes), width=Cm(5))

    # Right cell: document meta
    cell_meta = tbl.cell(0, 1)
    cell_meta.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    meta_lines = [
        ("Isotope Dashboard", True,  Pt(13), PURPLE),
        ("Technical Manual",  False, Pt(11), PURPLE),
        ("Internal — Confidential", False, Pt(9), GRAY_TEXT),
    ]
    for i, (txt, bold, sz, col) in enumerate(meta_lines):
        p = cell_meta.paragraphs[0] if i == 0 else cell_meta.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        _set_font(run, size=sz, bold=bold, color=col)

    # Thin purple rule below the header table
    p_rule = header.add_paragraph()
    pPr    = p_rule._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bot    = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "4B1E52")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Subsequent-page header (logo small, left-aligned)
    header2 = section.header
    for p in header2.paragraphs:
        p.clear()
    p2 = header2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run()
    run2.add_picture(io.BytesIO(logo_bytes), width=Cm(3))
    # Rule
    pPr2  = p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot2  = OxmlElement("w:bottom")
    bot2.set(qn("w:val"),   "single")
    bot2.set(qn("w:sz"),    "6")
    bot2.set(qn("w:space"), "1")
    bot2.set(qn("w:color"), "4B1E52")
    pBdr2.append(bot2)
    pPr2.append(pBdr2)


def _add_footer(doc):
    """Page number footer."""
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            p.clear()
        p_f = footer.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_f.add_run()
        _set_font(run_f, size=Pt(8), color=GRAY_TEXT)
        # Page number field
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run_f._r.append(fld)
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        run_f._r.append(instrText)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run_f._r.append(fld2)
        run_f2 = p_f.add_run("  |  Isotope Dashboard — Technical Manual  |  Confidential")
        _set_font(run_f2, size=Pt(8), color=GRAY_TEXT)

        # Purple top border on footer
        pPr  = p_f._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top  = OxmlElement("w:top")
        top.set(qn("w:val"),   "single")
        top.set(qn("w:sz"),    "6")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), "4B1E52")
        pBdr.append(top)
        pPr.append(pBdr)


# ── MANUAL.txt parser ─────────────────────────────────────────────────────────
# Line categories (applied in order):
#
#   TITLE     first line of the file
#   RULE      a line of === or --- or ─── or ═══ (separator, discard)
#   H1        e.g. "1. System Overview and Pipeline" (preceded by next line ===)
#   H2        e.g. "1.1  Purpose" (preceded by ---)
#   H3        e.g. "2.2a  Galliumbestralingen → raw.db: gallium_data"
#              (preceded by ─── or similar box-drawing)
#   BULLET    starts with "    - " or "    • "
#   CODE      starts with 4+ spaces and is not a bullet; or contains ┌ ─ etc.
#   BLANK     empty
#   TEXT      anything else

RULE_PATTERN   = re.compile(r'^[=\-─═━·]{4,}\s*$')
H1_PATTERN     = re.compile(r'^\d+\.\s+\S')         # "1. Something"
H2_PATTERN     = re.compile(r'^\d+\.\d+\s+\S')      # "1.1  Something"
H3_PATTERN     = re.compile(r'^\d+\.\d+[a-z]\s+\S') # "2.2a  Something"
BULLET_PATTERN = re.compile(r'^[ \t]{2,}[-•]\s')


def _classify_lines(lines: list[str]) -> list[tuple[str, str]]:
    """
    Return list of (category, text) for each content line.
    Requires one-look-ahead for underline detection.
    """
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Skip pure separator rules
        if RULE_PATTERN.match(stripped):
            i += 1
            continue

        # Look ahead for underline
        next_stripped = lines[i + 1].rstrip() if i + 1 < len(lines) else ""
        has_equals_ul = re.match(r'^[=]{4,}\s*$', next_stripped)
        has_dash_ul   = re.match(r'^[-]{4,}\s*$', next_stripped)
        has_box_ul    = re.match(r'^[─═━·]{4,}\s*$', next_stripped)

        if i == 0:
            result.append(("TITLE", stripped))
            i += 1
            continue

        if has_equals_ul:
            result.append(("H1", stripped))
            i += 2   # consume the underline
            continue

        if has_dash_ul:
            # Could be H1 (top-level content section) or H2 (sub-section)
            if H2_PATTERN.match(stripped):
                result.append(("H2", stripped))
            elif H1_PATTERN.match(stripped):
                result.append(("H1", stripped))
            else:
                result.append(("H2", stripped))
            i += 2
            continue

        if has_box_ul:
            result.append(("H3", stripped))
            i += 2
            continue

        if not stripped:
            result.append(("BLANK", ""))
            i += 1
            continue

        if BULLET_PATTERN.match(line):
            result.append(("BULLET", stripped.lstrip("-• ").strip()))
            i += 1
            continue

        # Indented non-bullet → code / preformatted
        if line.startswith("    ") and stripped:
            result.append(("CODE", stripped))
            i += 1
            continue

        result.append(("TEXT", stripped))
        i += 1

    return result


def _add_title_page(doc: Document, logo_bytes: bytes):
    """Insert a title page before the main content."""
    # Logo centred
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_logo, before=Pt(60), after=Pt(24))
    run_logo = p_logo.add_run()
    run_logo.add_picture(io.BytesIO(logo_bytes), width=Cm(8))

    # Main title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_title, before=Pt(24), after=Pt(8))
    r = p_title.add_run("Isotope Dashboard")
    _set_font(r, size=Pt(28), bold=True, color=PURPLE)

    # Subtitle
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_sub, before=Pt(0), after=Pt(4))
    r2 = p_sub.add_run("Technical Manual")
    _set_font(r2, size=Pt(18), bold=False, color=PINK)

    # Thin pink rule
    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_rule, before=Pt(12), after=Pt(12))
    pPr   = p_rule._p.get_or_add_pPr()
    pBdr  = OxmlElement("w:pBdr")
    bot   = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E50695")
    pBdr.append(bot)
    pPr.append(pBdr)

    # Meta lines
    meta = [
        "Cyclotron Production Monitoring System",
        "For developers and advanced users",
        "Curium Netherlands B.V. — Petten",
        "Confidential — internal use only",
    ]
    for m in meta:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(pm, before=Pt(2), after=Pt(2))
        rm = pm.add_run(m)
        _set_font(rm, size=Pt(10), color=GRAY_TEXT)

    # Page break
    doc.add_page_break()


def _render_lines(doc: Document, classified: list[tuple[str, str]]):
    """Convert classified lines into docx paragraphs."""
    in_code_block = False
    code_para     = None

    def _flush_code():
        nonlocal in_code_block, code_para
        in_code_block = False
        code_para     = None

    i = 0
    while i < len(classified):
        cat, text = classified[i]

        # ── H1 ─────────────────────────────────────────────────────────────
        if cat == "H1":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(18), after=Pt(4))
            r = p.add_run(text)
            _set_font(r, size=H1_SIZE, bold=True, color=PURPLE)
            # Purple underline rule
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "8")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "4B1E52")
            pBdr.append(bot)
            pPr.append(pBdr)

        # ── H2 ─────────────────────────────────────────────────────────────
        elif cat == "H2":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(12), after=Pt(3))
            r = p.add_run(text)
            _set_font(r, size=H2_SIZE, bold=True, color=PURPLE)

        # ── H3 ─────────────────────────────────────────────────────────────
        elif cat == "H3":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(8), after=Pt(2))
            r = p.add_run(text)
            _set_font(r, size=H3_SIZE, bold=True, color=PINK)

        # ── BULLET ─────────────────────────────────────────────────────────
        elif cat == "BULLET":
            _flush_code()
            p = doc.add_paragraph(style="List Bullet")
            _para_spacing(p, before=Pt(0), after=Pt(2))
            r = p.add_run(text)
            _set_font(r, size=BODY_SIZE)

        # ── CODE ────────────────────────────────────────────────────────────
        elif cat == "CODE":
            if not in_code_block:
                in_code_block = True
                # Start a new code paragraph
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(0), after=Pt(0), line=Pt(12))
            r = p.add_run(text)
            _set_font(r, name="Courier New", size=CODE_SIZE, color=RGBColor(0x33, 0x33, 0x33))
            # Grey background shading on the paragraph
            pPr  = p._p.get_or_add_pPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F2F2F2")
            pPr.append(shd)

        # ── BLANK ───────────────────────────────────────────────────────────
        elif cat == "BLANK":
            _flush_code()
            # Swallow consecutive blanks
            if i > 0 and classified[i - 1][0] != "BLANK":
                p = doc.add_paragraph()
                _para_spacing(p, before=Pt(0), after=Pt(4))

        # ── TEXT ────────────────────────────────────────────────────────────
        elif cat == "TEXT":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(0), after=Pt(4))
            r = p.add_run(text)
            _set_font(r, size=BODY_SIZE)

        # ── TITLE ───────────────────────────────────────────────────────────
        elif cat == "TITLE":
            pass   # already handled by title page

        i += 1


# ── Table of contents (simple manual list) ───────────────────────────────────
def _add_toc(doc: Document, classified: list[tuple[str, str]]):
    p_head = doc.add_paragraph()
    _para_spacing(p_head, before=Pt(0), after=Pt(8))
    r = p_head.add_run("Contents")
    _set_font(r, size=Pt(14), bold=True, color=PURPLE)

    for cat, text in classified:
        if cat == "H1":
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(2), after=Pt(1))
            r = p.add_run(text)
            _set_font(r, size=Pt(10), bold=True, color=PURPLE)
        elif cat == "H2":
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(0), after=Pt(0))
            # indent
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "360")
            pPr.append(ind)
            r = p.add_run(text)
            _set_font(r, size=Pt(9), color=BLACK)

    doc.add_page_break()


# ── Main ─────────────────────────────────────────────────────────────────────
def build():
    manual_path = Path(__file__).parent / "MANUAL.txt"
    out_path    = Path(__file__).parent / "Isotope_Dashboard_Manual.docx"

    lines      = manual_path.read_text(encoding="utf-8").splitlines()
    classified = _classify_lines(lines)

    # Build logo PNG in memory
    logo_bytes = _build_logo_png()

    doc = Document()
    _set_page_margins(doc)
    _add_header(doc, logo_bytes)
    _add_footer(doc)

    # Default Normal style
    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE

    # Title page
    _add_title_page(doc, logo_bytes)

    # Table of contents
    _add_toc(doc, classified)

    # Body
    _render_lines(doc, classified)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
