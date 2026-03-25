"""
build_manager_manual_docx.py
-----------------------------
Reads MANUAL_MANAGER_NL.txt and produces
Isotope_Dashboard_Managers_Handleiding.docx using the same Curium house
style as build_manual_docx.py (Calibri, brand colours, logo header).

Run:  python build_manager_manual_docx.py
"""

import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import docx.oxml

# ── Brand colours ────────────────────────────────────────────────────────────
PURPLE     = RGBColor(0x4B, 0x1E, 0x52)   # Curium dark purple
PINK       = RGBColor(0xE5, 0x06, 0x95)   # Curium hot pink
GRAY_TEXT  = RGBColor(0x80, 0x80, 0x80)   # Tagline / footer grey
BLACK      = RGBColor(0x00, 0x00, 0x00)
CODE_BG    = RGBColor(0xF2, 0xF2, 0xF2)

FONT_NAME  = "Calibri"
BODY_SIZE  = Pt(10)
H1_SIZE    = Pt(16)
H2_SIZE    = Pt(13)
H3_SIZE    = Pt(11)
CODE_SIZE  = Pt(9)


# ── Logo loader ───────────────────────────────────────────────────────────────
def _build_logo_png() -> bytes:
    from PIL import Image
    logo_path = Path(__file__).parent / "logo.png"
    img = Image.open(logo_path).convert("RGBA")
    bg  = Image.new("RGBA", img.size, (255, 255, 255, 255))
    bg.paste(img, mask=img.split()[3])
    rgb = bg.convert("RGB")
    buf = io.BytesIO()
    rgb.save(buf, format="PNG")
    return buf.getvalue()


# ── Drawing-id counter ────────────────────────────────────────────────────────
_DRAWING_ID = 0

def _next_drawing_id() -> int:
    global _DRAWING_ID
    _DRAWING_ID += 1
    return _DRAWING_ID


def _fix_drawing_ids(container):
    for el in container._element.iter(qn("wp:docPr")):
        el.set("id", str(_next_drawing_id()))


# ── docx helpers ──────────────────────────────────────────────────────────────
def _set_font(run, name=FONT_NAME, size=BODY_SIZE, bold=False,
              italic=False, color=BLACK, underline=False):
    run.font.name      = name
    run.font.size      = size
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.color.rgb = color
    run.font.underline = underline


def _para_spacing(para, before=Pt(0), after=Pt(4), line=None):
    pPr = para._p.get_or_add_pPr()
    spg = OxmlElement("w:spacing")
    spg.set(qn("w:before"), str(int(before.pt * 20)))
    spg.set(qn("w:after"),  str(int(after.pt  * 20)))
    if line is not None:
        spg.set(qn("w:line"),     str(int(line.pt * 20)))
        spg.set(qn("w:lineRule"), "exact")
    pPr.append(spg)


def _shade_cell(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)


def _set_page_margins(doc):
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.0)


def _add_header(doc, logo_bytes: bytes):
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    header  = section.first_page_header

    for p in header.paragraphs:
        p.clear()

    tbl = header.add_table(1, 2, width=Cm(16))
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl.style     = "Table Grid"
    tbl_elem = tbl._tbl
    tblPr    = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        tblBorders.append(border)
    tblPr.append(tblBorders)

    cell_logo = tbl.cell(0, 0)
    cell_logo.width = Cm(6)
    p_logo = cell_logo.paragraphs[0]
    run_logo = p_logo.add_run()
    run_logo.add_picture(io.BytesIO(logo_bytes), width=Cm(5))
    _fix_drawing_ids(header)

    cell_meta = tbl.cell(0, 1)
    cell_meta.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
    meta_lines = [
        ("Isotope Dashboard",         True,  Pt(13), PURPLE),
        ("Handleiding voor Managers", False, Pt(11), PURPLE),
        ("Intern — Vertrouwelijk",    False, Pt(9),  GRAY_TEXT),
    ]
    for i, (txt, bold, sz, col) in enumerate(meta_lines):
        p = cell_meta.paragraphs[0] if i == 0 else cell_meta.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(txt)
        _set_font(run, size=sz, bold=bold, color=col)

    p_rule = header.add_paragraph()
    pPr    = p_rule._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bot    = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "4B1E52")
    pBdr.append(bot)
    pPr.append(pBdr)

    header2 = section.header
    for p in header2.paragraphs:
        p.clear()
    p2 = header2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run()
    run2.add_picture(io.BytesIO(logo_bytes), width=Cm(3))
    _fix_drawing_ids(header2)
    pPr2  = p2._p.get_or_add_pPr()
    pBdr2 = OxmlElement("w:pBdr")
    bot2  = OxmlElement("w:bottom")
    bot2.set(qn("w:val"),   "single")
    bot2.set(qn("w:sz"),    "6")
    bot2.set(qn("w:space"), "1")
    bot2.set(qn("w:color"), "4B1E52")
    pBdr2.append(bot2)
    pPr2.append(pBdr2)


def _add_footer(doc):
    for section in doc.sections:
        footer = section.footer
        for p in footer.paragraphs:
            p.clear()
        p_f = footer.paragraphs[0]
        p_f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_f = p_f.add_run()
        _set_font(run_f, size=Pt(8), color=GRAY_TEXT)
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run_f._r.append(fld)
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        run_f._r.append(instrText)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run_f._r.append(fld2)
        run_f2 = p_f.add_run(
            "  |  Isotope Dashboard — Handleiding voor Managers"
            "  |  R.A.F. Sierkstra (SIE)  |  Vertrouwelijk"
        )
        _set_font(run_f2, size=Pt(8), color=GRAY_TEXT)

        pPr  = p_f._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top  = OxmlElement("w:top")
        top.set(qn("w:val"),   "single")
        top.set(qn("w:sz"),    "6")
        top.set(qn("w:space"), "4")
        top.set(qn("w:color"), "4B1E52")
        pBdr.append(top)
        pPr.append(pBdr)


# ── MANUAL_MANAGER_NL.txt parser ──────────────────────────────────────────────
RULE_PATTERN   = re.compile(r'^[=\-─═━·]{4,}\s*$')
H1_PATTERN     = re.compile(r'^\d+\.\s+\S')
H2_PATTERN     = re.compile(r'^\d+\.\d+\s+\S')
H3_PATTERN     = re.compile(r'^\d+\.\d+[a-z]\s+\S')
BULLET_PATTERN = re.compile(r'^[ \t]{2,}[-•]\s')
BOX_LINE_CHARS = set('┌│├└')
BOX_SEP_CHARS  = set('┌┬┐├┼┤└┴┘─━ ')


def _classify_lines(lines: list[str]) -> list[tuple[str, str]]:
    result = []
    i = 0
    while i < len(lines):
        line     = lines[i]
        stripped = line.rstrip()

        if RULE_PATTERN.match(stripped):
            i += 1
            continue

        next_stripped = lines[i + 1].rstrip() if i + 1 < len(lines) else ""
        has_equals_ul = re.match(r'^[=]{4,}\s*$', next_stripped)
        has_dash_ul   = re.match(r'^[-]{4,}\s*$', next_stripped)
        has_box_ul    = re.match(r'^[─═━·]{4,}\s*$', next_stripped)

        if i == 0:
            result.append(("TITLE", stripped))
            i += 1
            continue

        if has_equals_ul:
            result.append(("H1", stripped))
            i += 2
            continue

        if has_dash_ul:
            if H2_PATTERN.match(stripped):
                result.append(("H2", stripped))
            elif H1_PATTERN.match(stripped):
                result.append(("H1", stripped))
            else:
                result.append(("H2", stripped))
            i += 2
            continue

        if has_box_ul:
            result.append(("H3", stripped))
            i += 2
            continue

        if not stripped:
            result.append(("BLANK", ""))
            i += 1
            continue

        if BULLET_PATTERN.match(line):
            result.append(("BULLET", stripped.lstrip("-• ").strip()))
            i += 1
            continue

        if line.startswith("    ") and stripped:
            result.append(("CODE", stripped))
            i += 1
            continue

        if stripped and stripped[0] in BOX_LINE_CHARS:
            result.append(("BOXTABLE", stripped))
            i += 1
            continue

        result.append(("TEXT", stripped))
        i += 1

    return result


def _add_title_page(doc: Document, logo_bytes: bytes):
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_logo, before=Pt(60), after=Pt(24))
    run_logo = p_logo.add_run()
    run_logo.add_picture(io.BytesIO(logo_bytes), width=Cm(8))
    _fix_drawing_ids(doc)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_title, before=Pt(24), after=Pt(8))
    r = p_title.add_run("Isotope Dashboard")
    _set_font(r, size=Pt(28), bold=True, color=PURPLE)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_sub, before=Pt(0), after=Pt(4))
    r2 = p_sub.add_run("Handleiding voor Managers")
    _set_font(r2, size=Pt(18), bold=False, color=PINK)

    p_rule = doc.add_paragraph()
    p_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _para_spacing(p_rule, before=Pt(12), after=Pt(12))
    pPr  = p_rule._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "12")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "E50695")
    pBdr.append(bot)
    pPr.append(pBdr)

    meta = [
        "Cyclotron Productiemonitoring Systeem",
        "Voor managers",
        "Curium Netherlands B.V. — Petten",
        "Intern — Vertrouwelijk",
    ]
    for m in meta:
        pm = doc.add_paragraph()
        pm.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _para_spacing(pm, before=Pt(2), after=Pt(2))
        rm = pm.add_run(m)
        _set_font(rm, size=Pt(10), color=GRAY_TEXT)

    doc.add_paragraph()
    atbl = doc.add_table(rows=3, cols=2)
    atbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    atbl_elem = atbl._tbl
    atPr = atbl_elem.find(qn("w:tblPr"))
    if atPr is None:
        atPr = OxmlElement("w:tblPr")
        atbl_elem.insert(0, atPr)
    atBdr = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none")
        atBdr.append(b)
    atPr.append(atBdr)

    rows_data = [
        ("Auteur",    "R.A.F. Sierkstra  (SIE)"),
        ("Functie",   "Process Reliability Engineer"),
        ("Contact",   "ext. 7062  ·  reinder.sierkstra@curiumpharma.com"),
    ]
    for i, (label, value) in enumerate(rows_data):
        lbl_cell = atbl.cell(i, 0)
        val_cell = atbl.cell(i, 1)
        lbl_cell._tc.get_or_add_tcPr().append(OxmlElement("w:tcW"))
        lbl_cell._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:w"), "1200")
        lbl_cell._tc.get_or_add_tcPr().find(qn("w:tcW")).set(qn("w:type"), "dxa")

        lp = lbl_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _para_spacing(lp, before=Pt(1), after=Pt(1))
        rl = lp.add_run(label)
        _set_font(rl, size=Pt(9), bold=True, color=PURPLE)

        vp = val_cell.paragraphs[0]
        _para_spacing(vp, before=Pt(1), after=Pt(1))
        rv = vp.add_run("  " + value)
        _set_font(rv, size=Pt(9), color=GRAY_TEXT)

    doc.add_page_break()


def _parse_boxtable(raw_lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in raw_lines:
        if all(c in BOX_SEP_CHARS for c in line):
            continue
        parts = line.split('│')
        if len(parts) < 3:
            continue
        raw_cells = parts[1:-1]
        cells     = [p.strip() for p in raw_cells]
        if not cells:
            continue
        is_continuation = (
            rows and
            (not cells[0] or (len(raw_cells[0]) > 1 and raw_cells[0].startswith('  ')))
        )
        if is_continuation:
            for j, text in enumerate(cells):
                if text and j < len(rows[-1]):
                    sep = '' if j == 0 else ' '
                    rows[-1][j] = (rows[-1][j] + sep + text).strip()
        else:
            rows.append(cells)
    return rows


def _render_boxtable(doc: Document, raw_lines: list[str]):
    rows = _parse_boxtable(raw_lines)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl   = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    _para_spacing(doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph(),
                  before=Pt(6), after=Pt(6))
    PAGE_W = Cm(16)
    col0_w = Cm(4.5) if ncols > 1 else PAGE_W
    rest_w = (PAGE_W - col0_w) // max(ncols - 1, 1)
    for r_idx, row_cells in enumerate(rows):
        for c_idx, cell_text in enumerate(row_cells):
            cell = tbl.cell(r_idx, c_idx)
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcW  = OxmlElement("w:tcW")
            width_emu = col0_w if c_idx == 0 else rest_w
            twips = int(width_emu * 1440 / 914400)
            tcW.set(qn("w:w"),    str(twips))
            tcW.set(qn("w:type"), "dxa")
            tcPr.append(tcW)
            p   = cell.paragraphs[0]
            _para_spacing(p, before=Pt(2), after=Pt(2))
            run = p.add_run(cell_text)
            if c_idx == 0:
                _set_font(run, size=Pt(9), bold=True, color=PURPLE)
            else:
                _set_font(run, size=Pt(9), color=BLACK)
    for r_idx in range(len(rows)):
        _shade_cell(tbl.cell(r_idx, 0), "EDE7F0")
    doc.add_paragraph()


def _render_lines(doc: Document, classified: list[tuple[str, str]]):
    in_code_block = False

    def _flush_code():
        nonlocal in_code_block
        in_code_block = False

    i = 0
    while i < len(classified):
        cat, text = classified[i]

        if cat == "H1":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(18), after=Pt(4))
            r = p.add_run(text)
            _set_font(r, size=H1_SIZE, bold=True, color=PURPLE)
            pPr  = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "8")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), "4B1E52")
            pBdr.append(bot)
            pPr.append(pBdr)

        elif cat == "H2":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(12), after=Pt(3))
            r = p.add_run(text)
            _set_font(r, size=H2_SIZE, bold=True, color=PURPLE)

        elif cat == "H3":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(8), after=Pt(2))
            r = p.add_run(text)
            _set_font(r, size=H3_SIZE, bold=True, color=PINK)

        elif cat == "BULLET":
            _flush_code()
            p = doc.add_paragraph(style="List Bullet")
            _para_spacing(p, before=Pt(0), after=Pt(2))
            r = p.add_run(text)
            _set_font(r, size=BODY_SIZE)

        elif cat == "CODE":
            if not in_code_block:
                in_code_block = True
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(0), after=Pt(0), line=Pt(12))
            r = p.add_run(text)
            _set_font(r, name="Courier New", size=CODE_SIZE,
                      color=RGBColor(0x33, 0x33, 0x33))
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "F2F2F2")
            pPr.append(shd)

        elif cat == "BOXTABLE":
            _flush_code()
            raw_lines = [text]
            while i + 1 < len(classified) and classified[i + 1][0] == "BOXTABLE":
                i += 1
                raw_lines.append(classified[i][1])
            _render_boxtable(doc, raw_lines)

        elif cat == "BLANK":
            _flush_code()
            if i > 0 and classified[i - 1][0] != "BLANK":
                p = doc.add_paragraph()
                _para_spacing(p, before=Pt(0), after=Pt(4))

        elif cat == "TEXT":
            _flush_code()
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(0), after=Pt(4))
            r = p.add_run(text)
            _set_font(r, size=BODY_SIZE)

        elif cat == "TITLE":
            pass   # handled by title page

        i += 1


def _add_toc(doc: Document, classified: list[tuple[str, str]]):
    p_head = doc.add_paragraph()
    _para_spacing(p_head, before=Pt(0), after=Pt(8))
    r = p_head.add_run("Inhoud")
    _set_font(r, size=Pt(14), bold=True, color=PURPLE)

    for cat, text in classified:
        if cat == "H1":
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(2), after=Pt(1))
            r = p.add_run(text)
            _set_font(r, size=Pt(10), bold=True, color=PURPLE)
        elif cat == "H2":
            p = doc.add_paragraph()
            _para_spacing(p, before=Pt(0), after=Pt(0))
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "360")
            pPr.append(ind)
            r = p.add_run(text)
            _set_font(r, size=Pt(9), color=BLACK)

    doc.add_page_break()


# ── Main ─────────────────────────────────────────────────────────────────────
def build():
    manual_path = Path(__file__).parent / "MANUAL_MANAGER_NL.txt"
    out_path    = Path(__file__).parent / "Isotope_Dashboard_Managers_Handleiding.docx"

    lines      = manual_path.read_text(encoding="utf-8").splitlines()
    classified = _classify_lines(lines)

    logo_bytes = _build_logo_png()

    doc = Document()
    _set_page_margins(doc)
    _add_header(doc, logo_bytes)
    _add_footer(doc)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE

    _add_title_page(doc, logo_bytes)
    _add_toc(doc, classified)
    _render_lines(doc, classified)

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build()
